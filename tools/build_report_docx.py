"""
Build Word report from eval JSON + optional visualization folder.
Requires: pip install python-docx
"""

import argparse
import json
import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def add_title(doc: Document, text: str) -> None:
    t = doc.add_heading(text, level=0)  # 文档主标题（标题样式）
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中排版，便于毕设/答辩材料直接使用


def main():
    parser = argparse.ArgumentParser("Build RIS results Word report")
    parser.add_argument("--eval-json", type=str, required=True)
    parser.add_argument("--vis-dir", type=str, default="", help="Folder with *_pred_overlay.jpg etc.")
    parser.add_argument("--out-docx", type=str, required=True)
    parser.add_argument("--embed-images", type=int, default=6, help="Max overlay images to embed")
    args = parser.parse_args()

    with open(args.eval_json, "r", encoding="utf-8") as f:
        results = json.load(f)

    doc = Document()  # 空 Word 文档对象
    style = doc.styles["Normal"]  # 正文样式
    style.font.name = "宋体"  # 中文字体（Windows 常见）
    style.font.size = Pt(10.5)  # 五号左右，符合一般课程报告习惯

    add_title(doc, "基于文本引导的图像语义分割实验成果说明")
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    doc.add_heading("一、任务与工程说明", level=1)
    doc.add_paragraph(
        "本实验在 RefCOCO 标注与 COCO train2014 图像上，使用 PyTorch 与 CLIP 文本编码器实现指代表达分割（"
        "Referring Image Segmentation, RIS）实验系统。训练采用 BCE + Dice 损失，验证与测试统计 mean IoU "
        "（逐 batch 均值，与 train.py 中验证口径一致）。"
    )

    doc.add_heading("二、评估流程", level=1)
    doc.add_paragraph(
        "1. 使用 tools/build_refcoco_index.py 将官方 RefCOCO 转为索引 JSON 与掩码；"
        "2. 使用 train.py 训练并保存 checkpoints/best.pt；"
        "3. 使用 eval.py 在 val / testA / testB 上评估；"
        "4. 使用 visualize.py 抽样生成叠加可视化图。"
    )
    p = doc.add_paragraph("评估命令示例（已在结果 JSON 中记录检查点路径）：")
    p = doc.add_paragraph(style="List Bullet")
    run = (
        'python eval.py --data-root "<refcoco_ready>" --checkpoint "<best.pt>" '
        '--split val=splits/val.json --split testA=splits/testA.json --split testB=splits/testB.json '
        "--save-json eval_results.json"
    )
    doc.add_paragraph(run)

    doc.add_heading("三、定量结果", level=1)
    ckpt = results.get("checkpoint", "")
    doc.add_paragraph(f"检查点文件：{ckpt}")
    if results.get("ckpt_epoch") is not None:
        doc.add_paragraph(f"检查点记录的 epoch：{results.get('ckpt_epoch')}")
    best_miou = results.get("ckpt_best_val_miou", results.get("ckpt_best_val_iou"))
    if best_miou is not None:
        doc.add_paragraph(f"检查点记录的 best val mIoU：{best_miou:.4f}")

    table = doc.add_table(rows=1, cols=5)  # 划分、样本数、mIoU、cIoU、损失
    hdr = table.rows[0].cells
    hdr[0].text = "划分"
    hdr[1].text = "样本数"
    hdr[2].text = "mIoU"
    hdr[3].text = "cIoU"
    hdr[4].text = "Mean Loss"

    splits = results.get("splits", {})
    preferred = ["val", "testA", "testB", "train"]
    names = [k for k in preferred if k in splits] + sorted(k for k in splits if k not in preferred)
    for name in names:
        row = table.add_row().cells
        s = splits[name]
        miou = s.get("mIoU", s.get("mean_iou", 0))
        ciou = s.get("cIoU", 0)
        row[0].text = name
        row[1].text = str(s.get("num_samples", ""))
        row[2].text = f"{miou:.4f}"
        row[3].text = f"{ciou:.4f}"
        row[4].text = f"{s.get('mean_loss', 0):.4f}"

    doc.add_paragraph()

    doc.add_heading("四、可视化流程与成果", level=1)
    doc.add_paragraph(
        "可视化脚本对索引中随机抽样的图像生成：原图（*_input.jpg）、预测叠加（绿色，*_pred_overlay.jpg）、"
        "真值叠加（红色，*_gt_overlay.jpg），以及对应文本（*_text.txt）。IoU 在原始分辨率上根据二值掩码计算。"
    )

    if args.vis_dir and os.path.isdir(args.vis_dir):
        doc.add_paragraph(f"可视化输出目录：{os.path.abspath(args.vis_dir)}")
        pat = [f for f in os.listdir(args.vis_dir) if f.endswith("_pred_overlay.jpg")]
        pat.sort()
        doc.add_paragraph(f"共生成预测叠加图 {len(pat)} 张（本节展示至多 {args.embed_images} 张）。")
        for fn in pat[: args.embed_images]:  # 控制 docx 体积，避免插图过多
            path = os.path.join(args.vis_dir, fn)
            doc.add_paragraph(fn, style="Caption")  # 图片题注（文件名含 IoU 提示）
            try:
                doc.add_picture(path, width=Inches(5.5))  # 统一宽度，A4 页内可读
            except Exception as e:
                doc.add_paragraph(f"（插图失败：{e}）")
    else:
        doc.add_paragraph("未提供可视化目录或目录不存在，可运行 visualize.py 后重新生成本文档。")

    doc.add_heading("五、说明与后续工作", level=1)
    doc.add_paragraph(
        "本报告中的 IoU 为工程内一致性指标；若与论文表格严格对齐，需确认是否与官方实现使用相同的 masks、"
        "阈值与平均方式。后续可替换更强分割主干、解冻或微调 CLIP、以及增加指代特定的注意力模块以提升精度。"
    )

    out = os.path.abspath(args.out_docx)
    os.makedirs(os.path.dirname(out) or ".", exist_ok=True)
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
