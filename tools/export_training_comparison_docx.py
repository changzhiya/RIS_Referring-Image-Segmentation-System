"""
Scan ris_mvp/result/**/val_metrics.jsonl and write a Word comparison report.
Requires: pip install python-docx
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:
    print("Install: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def load_metrics(path: Path) -> list[dict]:
    rows: list[dict] = []
    with open(path, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                rows.append(json.loads(line))
            except json.JSONDecodeError:
                continue
    return rows


def summarize(rows: list[dict]) -> dict | None:
    if not rows:
        return None
    best_miou = max(rows, key=lambda r: float(r.get("val_mIoU", 0)))
    best_ciou = max(rows, key=lambda r: float(r.get("val_cIoU", 0)))
    last = rows[-1]
    first = rows[0]
    min_loss = min(rows, key=lambda r: float(r.get("val_loss", 999)))
    return {
        "n": len(rows),
        "best_miou": float(best_miou["val_mIoU"]),
        "best_miou_ep": int(best_miou["epoch"]),
        "best_ciou": float(best_ciou["val_cIoU"]),
        "best_ciou_ep": int(best_ciou["epoch"]),
        "last_miou": float(last["val_mIoU"]),
        "last_ciou": float(last["val_cIoU"]),
        "last_loss": float(last["val_loss"]),
        "first_miou": float(first["val_mIoU"]),
        "first_ciou": float(first["val_cIoU"]),
        "min_loss": float(min_loss["val_loss"]),
        "min_loss_ep": int(min_loss["epoch"]),
    }


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(9)


def main() -> None:
    script_dir = Path(__file__).resolve().parent
    result_root = script_dir.parent / "result"
    # ASCII filename avoids mojibake on some Windows shells; title inside doc is Chinese.
    out_path = result_root / "training_experiments_comparison.docx"

    runs: list[tuple[str, Path, list[dict], dict | None]] = []
    for p in sorted(result_root.glob("**/val_metrics.jsonl")):
        rows = load_metrics(p)
        runs.append((p.parent.name, p, rows, summarize(rows)))

    doc = Document()
    title = doc.add_heading("RefCOCO RIS 训练实验对比报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(
        "数据来源：各实验目录下 val_metrics.jsonl（每行一个 epoch 的验证集均值指标）。"
    )
    p.add_run(
        " val_loss 为加权 BCE+Dice+softIoU，不同实验损失尺度不可横向比较；对比以 val_mIoU / val_cIoU 为主。"
    )

    doc.add_heading("1. 实验汇总对比表", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = [
        "实验目录",
        "记录条数",
        "最佳 val_mIoU",
        "对应 epoch",
        "最佳 val_cIoU",
        "cIoU 对应 epoch",
        "末行 val_mIoU / cIoU",
    ]
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h)

    for folder, path, rows, s in runs:
        if s is None:
            row = table.add_row().cells
            set_cell_text(row[0], folder)
            set_cell_text(row[1], "0（无数据）")
            for j in range(2, 7):
                set_cell_text(row[j], "-")
            continue
        row = table.add_row().cells
        set_cell_text(row[0], folder)
        set_cell_text(row[1], str(s["n"]))
        set_cell_text(row[2], f"{s['best_miou']:.4f}")
        set_cell_text(row[3], str(s["best_miou_ep"]))
        set_cell_text(row[4], f"{s['best_ciou']:.4f}")
        set_cell_text(row[5], str(s["best_ciou_ep"]))
        set_cell_text(row[6], f"{s['last_miou']:.4f} / {s['last_ciou']:.4f}")

    doc.add_heading("2. 各实验配置与结论摘要", level=1)
    notes = [
        (
            "checkpoints_doc31",
            "文档 3.1.x 预设（--doc-stage1）：224 输入、BCE+Dice+softIoU 0.5/0.3/0.2、CLIP 文本末 2 层解冻等。",
        ),
        (
            "checkpoints_miou10_20260414_215059",
            "验证 IoU 预设（--preset-early-val-miou）：256 输入、损失 0.15/0.45/0.40、CLIP 末 1 层、略缓 RandomResizedCrop 等；本表为完整 14 epoch 曲线。",
        ),
        (
            "checkpoints_active",
            "历史默认/活跃目录实验：256、较长训练（至 14 epoch）；损失权重与 doc 预设不同，val_loss 数值偏大属正常。",
        ),
        (
            "checkpoints_retrain",
            "retrain 目录：条数较少且存在重复 epoch 行，可能为中断/重复写入；解读时以同 epoch 最后一次或人工去重为准。",
        ),
    ]
    for name, text in notes:
        doc.add_paragraph(f"{name}：{text}", style="List Bullet")

    doc.add_heading("3. 简要对比结论", level=1)
    valid = [(f, r, s) for f, _, r, s in runs if s is not None and r]
    if valid:
        by_miou = max(valid, key=lambda x: x[2]["best_miou"])
        by_ciou = max(valid, key=lambda x: x[2]["best_ciou"])
        sm, sc = by_miou[2], by_ciou[2]
        doc.add_paragraph(
            f"按本目录 jsonl 统计：val_mIoU 最高为「{by_miou[0]}」（{sm['best_miou']:.4f}，epoch {sm['best_miou_ep']}）；"
            f"val_cIoU 最高为「{by_ciou[0]}」（{sc['best_ciou']:.4f}，epoch {sc['best_ciou_ep']}）。"
        )
        doc.add_paragraph(
            "doc31 与 miou10 预设输入分辨率与损失不同，不宜仅比较 val_loss；论文/报告建议并列给出 mIoU、cIoU 与训练设置表。"
        )

    doc.add_heading("4. 各实验逐 epoch 指标（节选）", level=1)
    for folder, path, rows, s in runs:
        if not rows:
            continue
        doc.add_heading(f"4.x {folder}", level=2)
        doc.add_paragraph(str(path))
        t2 = doc.add_table(rows=1, cols=6)
        t2.style = "Table Grid"
        h2 = t2.rows[0].cells
        for i, h in enumerate(["epoch", "val_loss", "val_mIoU", "val_cIoU", "train_loss", "train_iou"]):
            set_cell_text(h2[i], h)
        for r in rows[:20]:
            rr = t2.add_row().cells
            set_cell_text(rr[0], str(int(r["epoch"])))
            set_cell_text(rr[1], f"{float(r['val_loss']):.4f}")
            set_cell_text(rr[2], f"{float(r['val_mIoU']):.4f}")
            set_cell_text(rr[3], f"{float(r['val_cIoU']):.4f}")
            set_cell_text(rr[4], f"{float(r['train_loss']):.4f}")
            set_cell_text(rr[5], f"{float(r['train_iou_batch_mean']):.4f}")
        if len(rows) > 20:
            doc.add_paragraph(f"（仅列前 20 行，共 {len(rows)} 行；完整数据见源 jsonl）")

    doc.save(out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()
